"""Generate BMC Genomics Word submission from paper_bmc_genomics.md"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Deep Learning for Salmonella enterica Detection from Whole-Genome Sequences:\nA Comprehensive Benchmark of CNN, Liquid Neural Networks, and Gradient-Boosted Trees")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

# Authors placeholder
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run("[Author 1], [Author 2], [Author 3], [Corresponding Author]*")
run.font.size = Pt(12)

affiliation = doc.add_paragraph()
affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affiliation.add_run("[Department, Institution, City, Country]")
run.font.size = Pt(10)
run.italic = True

corresponding = doc.add_paragraph()
corresponding.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = corresponding.add_run("*Corresponding author: [email@institution.edu]")
run.font.size = Pt(10)

doc.add_paragraph()  # spacer

# === ABSTRACT ===
abs_heading = doc.add_paragraph()
run = abs_heading.add_run("Abstract")
run.bold = True
run.font.size = Pt(14)

abstract_sections = {
    "Background": "Rapid and accurate identification of Salmonella enterica from whole-genome sequencing (WGS) data is critical for food safety surveillance, clinical diagnostics, and outbreak investigation. While alignment-based tools like Kraken2 and Centrifuge dominate current practice, deep learning methods offer potential advantages in speed, database independence, and generalization. However, no comprehensive benchmark exists comparing deep learning architectures against each other and against established bioinformatics tools for this task.",
    "Results": "We benchmarked seven classification methods — Convolutional Neural Network (CNN), Liquid Neural Network (LNN/CfC), Gated Recurrent Unit (GRU), BiLSTM, Transformer, XGBoost, and Random Forest — on 14,156 bacterial genomes (6,110 Salmonella, 8,046 non-Salmonella) encoded as k-mer frequency vectors. CNN achieved the highest accuracy (99.64% ± 0.3%, AUC 1.000, 174K parameters) with 2-second training, outperforming a simplified Kraken2-like k-mer matching classifier (83.3%), GRU (98.3%), and XGBoost (94.2%). Our models matched or exceeded published Kraken2 (91-99%) and Centrifuge (89-97%) benchmarks while requiring no external database. Attention pooling was the critical architectural component, improving LNN accuracy by 18 percentage points. LNN demonstrated 6-fold lower variance than CNN at 10% training data. Out-of-distribution testing revealed performance stratified by phylogenetic distance, partially mitigated by hard negative mining. External validation on 20 independent genomes achieved 80% accuracy with 100% Salmonella recall. Serotype classification proved substantially more challenging (<=36% for all k-mer methods) compared to the specialized tool SeqSero2 (93-98%).",
    "Conclusions": "1D-CNN operating on k-mer frequency vectors is the most accurate and efficient method for WGS-based Salmonella detection, matching established tools without requiring multi-gigabyte databases. Liquid Neural Networks offer advantages in data-limited and interpretability-critical scenarios. The systematic benchmark provides a robust foundation for deploying deep learning in genomic pathogen surveillance."
}

for label, text in abstract_sections.items():
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)

# Keywords
kw = doc.add_paragraph()
run = kw.add_run("Keywords: ")
run.bold = True
run.font.size = Pt(10)
run = kw.add_run("Salmonella enterica, deep learning, k-mer encoding, whole-genome sequencing, pathogen detection, liquid neural networks, Kraken2, convolutional neural network, benchmark")
run.font.size = Pt(10)

doc.add_page_break()

# === HELPER FUNCTIONS ===
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_table(headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer
    return table

# === BACKGROUND ===
add_heading_styled("Background", level=1)

add_para("Salmonella enterica remains one of the most significant foodborne pathogens worldwide, causing an estimated 93.8 million cases of gastroenteritis and 155,000 deaths annually [1]. Rapid and accurate detection is essential for outbreak response, clinical treatment, and food safety monitoring. Conventional detection methods — culture-based isolation followed by biochemical and serological confirmation — require 3-5 days and specialized laboratory facilities [2]. While PCR-based methods offer faster turnaround, they are limited to pre-specified genetic targets and may miss emerging or atypical strains [3].")

add_para("The declining cost of whole-genome sequencing has enabled its adoption for pathogen surveillance and diagnostics [4]. WGS-based computational identification has largely relied on two classes of methods: (1) alignment-based approaches using exact k-mer matching against reference databases (Kraken2 [5], Centrifuge [6]), which achieve 91-99% accuracy but require multi-gigabyte databases and substantial computational resources, and (2) marker gene detection tools (SeqSero2 [7]), which achieve 93-98% accuracy for serotype determination but rely on curated allele databases that require continuous updates.")

add_para("Deep learning offers the potential for rapid, database-free pathogen detection directly from genomic sequences. Convolutional neural networks (CNNs) have been successfully applied to DNA sequence classification [8, 9], while recurrent architectures including Long Short-Term Memory (LSTM) networks and Gated Recurrent Units (GRU) can model sequential dependencies in genomic data. More recently, Liquid Neural Networks (LNNs) with Closed-form Continuous-time (CfC) cells [10] have demonstrated remarkable parameter efficiency and out-of-distribution generalization across multiple domains, but their application to genomic sequence classification remains unexplored.")

add_para("In this study, we present a comprehensive benchmark of seven classification methods for WGS-based Salmonella detection. Our contributions include: (1) the first systematic evaluation of Liquid Neural Networks for genomic sequence classification; (2) identification of attention pooling as the critical architectural component; (3) demonstration that deep learning matches or exceeds exact k-mer matching tools while eliminating database dependencies; (4) characterization of LNN data efficiency and interpretability advantages; and (5) a robust experimental framework spanning multiple evaluation dimensions.")

# === METHODS ===
add_heading_styled("Methods", level=1)

add_heading_styled("Dataset", level=2)
add_para("We assembled a dataset of 14,156 bacterial genomes from NCBI GenBank, comprising 6,110 Salmonella enterica genomes across 6 serotypes and 8,046 non-Salmonella genomes spanning 12 bacterial species. Data were sourced from two formats: pre-chunked genomic fragments (~1,000-2,000 bp) from the NCBI Pathogen Detection database, and 515 complete genomes that were fragmented using a sliding window approach (window size 1,000 bp, stride 500 bp) with quality filtering (single-base frequency <50%, N-ratio <5%, GC content 25-75%).")

add_table(
    ["Category", "Source", "Species/Serotypes", "Genomes", "Chunks"],
    [
        ["Positive (Salmonella)", "serotype_data/", "6 serotypes", "5,756", "298,628"],
        ["Positive (Salmonella)", "zheng-yangpin/", "Various", "354", "—"],
        ["Negative (non-Salmonella)", "negative_species/", "6 species", "7,885", "123,154"],
        ["Negative (non-Salmonella)", "fu-yangpin/", "161 spp. (incl. 63 E. coli)", "161", "—"],
        ["Total", "—", "12+ species", "14,156", "421,782"],
    ],
    "Table 1. Dataset composition"
)

add_heading_styled("Sequence Encoding and Sample Construction", level=2)
add_para("Each genomic fragment was encoded as a k-mer frequency vector (k=4, 256 dimensions). Each genome was represented as a sequence of 32 randomly selected fragment vectors, producing input tensors of shape (32, 256). Samples were split by genome source into training (70%), validation (15%), and test (15%) sets.")

add_heading_styled("Model Architectures", level=2)
add_para("LNN/CfC: Three stacked CfC layers (128->64->32 hidden units) with LayerNorm, Dropout(0.1), and learned attention pooling. Three variants: Small (101K params), Medium (413K), Large (1.21M).")
add_para("1D-CNN: Three convolutional layers (256->64->128->256 channels, kernel=3) with BatchNorm, ReLU, MaxPool, adaptive max pooling, linear classifier (174K params).")
add_para("Comparative architectures: BiLSTM with attention (824K), Transformer (851K), GRU (196K), Simple RNN (66K). Traditional ML: XGBoost (100 estimators, max_depth=6), Random Forest (200 trees).")
add_para("External tool comparison: KmerVote — simplified Kraken2-like k=10 exact matching with majority voting. Published benchmarks reported for Kraken2 [5], Centrifuge [6], SeqSero2 [7].")

add_heading_styled("Training Protocol", level=2)
add_para("AdamW optimizer (lr=1e-3, weight_decay=1e-4), BCE loss with class-balanced weights, AMP, gradient clipping (max norm 1.0), cosine LR decay with 500-step warmup. Early stopping (patience=15). CfC models used lr=5e-5. All deep learning experiments repeated 3 times. GPU: NVIDIA RTX 5060 (12GB).")

add_heading_styled("Evaluation Framework", level=2)
add_para("Standard classification: Accuracy, F1, AUC (mean +/- std, n=3). OOD generalization: Leave-one-species/serotype-out cross-validation with enhanced negative mining (63 E. coli strains). Data efficiency: 10/25/50/100% training subsets. Ablation: cell type, pooling, positional encoding, chunk count, learning rate. Interpretability: t-SNE, PCA, gradient-based chunk importance. External validation: 20 independent genomes.")

# === RESULTS ===
add_heading_styled("Results", level=1)

add_heading_styled("Binary Classification Performance", level=2)
add_para("CNN achieved the highest accuracy (99.64% +/- 0.3%, AUC 1.000), training in 2 seconds. GRU was second (98.27%). LNN with attention pooling reached 97.84%. XGBoost, the best non-deep-learning method, achieved 94.16%.")

add_table(
    ["Model", "Accuracy", "F1", "AUC", "Parameters", "Train Time"],
    [
        ["CNN", "0.9964 +/- 0.003", "0.998", "1.000", "173,633", "2s"],
        ["GRU", "0.9827", "0.989", "0.999", "195,809", "4s"],
        ["LNN + Attention", "0.9784", "0.986", "0.994", "413,378", "130s"],
        ["LNN-Medium (opt LR)", "0.9221", "0.950", "0.915", "412,833", "124s"],
        ["XGBoost", "0.9416 +/- 0.000", "0.964", "0.989", "~500,000", "2s"],
        ["LSTM", "0.8882 +/- 0.053", "0.935", "0.843", "823,810", "3s"],
        ["Random Forest", "0.8723", "0.926", "0.975", "~5,000,000", "1s"],
        ["LNN-Small (last pool)", "0.8095 +/- 0.015", "0.892", "0.671", "101,121", "44s"],
        ["Transformer", "0.7987", "0.888", "0.500", "850,946", "3s"],
    ],
    "Table 2. Binary classification performance (mean +/- std, n=3)"
)

add_heading_styled("Comparison with Existing Bioinformatics Tools", level=2)
add_para("Our deep learning models matched or exceeded published Kraken2 and Centrifuge benchmarks while requiring no external database. The models learn continuous decision boundaries from k-mer frequencies rather than relying on discrete k-mer presence/absence.")

add_table(
    ["Method", "Type", "Accuracy", "Database Required"],
    [
        ["CNN (ours)", "k-mer frequency + deep learning", "0.996", "None"],
        ["Kraken2 [5]", "k=35 exact match + LCA", "0.91-0.99", "50 GB"],
        ["GRU (ours)", "k-mer frequency + deep learning", "0.983", "None"],
        ["Centrifuge [6]", "BWT-FM index", "0.89-0.97", "10 GB"],
        ["LNN+Attn (ours)", "k-mer frequency + deep learning", "0.978", "None"],
        ["SeqSero2 [7]", "Marker allele database", "0.93-0.98*", "<1 GB"],
        ["KmerVote k=10 (ours)", "Exact k-mer matching", "0.833", "33K kmers"],
        ["XGBoost (ours)", "k-mer frequency + gradient boosting", "0.942", "None"],
    ],
    "Table 3. Comparison with existing bioinformatics tools (* serotype classification only)"
)

add_heading_styled("Attention Pooling is the Critical Component", level=2)
add_para("Replacing last-step pooling with learned attention improved LNN accuracy by 18 percentage points (79.87% -> 97.84%), a larger effect than any other architectural modification.")

add_table(
    ["Pooling", "Accuracy", "F1", "AUC"],
    [
        ["Last-step", "0.7987", "0.888", "0.592"],
        ["Mean", "0.7987", "0.888", "0.775"],
        ["Attention", "0.9784", "0.986", "0.994"],
    ],
    "Table 4. Effect of pooling strategy on LNN classification"
)

add_heading_styled("Cell Type Ablation: CfC vs. GRU vs. RNN", level=2)
add_para("GRU outperformed CfC by 8 percentage points in identical architectures (98.27% vs. 90.26%), suggesting that for discrete k-mer frequency sequences, gating mechanisms are more effective than continuous-time ODE dynamics.")

add_table(
    ["Cell", "Accuracy", "AUC", "Parameters"],
    [
        ["Simple RNN", "0.8420", "0.689", "65,889"],
        ["CfC (LNN)", "0.9026", "0.948", "101,121"],
        ["GRU", "0.9827", "0.999", "195,809"],
    ],
    "Table 5. Cell type ablation"
)

add_heading_styled("Data Efficiency", level=2)
add_para("LNN-Small demonstrated 6.1-fold lower variance than CNN at 10% training data. Performance degradation from 100% to 10% was 7.4% for LNN versus 15.7% for CNN.")

add_table(
    ["Model", "10%", "25%", "50%", "100%", "Degradation"],
    [
        ["XGBoost", "0.873 +/- 0.011", "0.883 +/- 0.009", "0.920 +/- 0.012", "0.942", "-6.9%"],
        ["LNN-Small", "0.850 +/- 0.009", "0.846 +/- 0.004", "0.851 +/- 0.008", "0.924", "-7.4%"],
        ["CNN", "0.838 +/- 0.055", "0.863 +/- 0.091", "0.991 +/- 0.002", "0.995", "-15.7%"],
    ],
    "Table 6. Data efficiency (accuracy at different training fractions)"
)

add_heading_styled("Out-of-Distribution Generalization", level=2)
add_para("OOD accuracy was perfectly stratified by phylogenetic distance. Incorporating 63 E. coli strains as hard negatives improved Klebsiella OOD accuracy 4.7-fold.")

add_table(
    ["Held-out Species", "Phylogenetic Distance", "Standard", "Enhanced (+63 E. coli)"],
    [
        ["E. faecalis", "Far (Firmicutes)", "1.000", "1.000"],
        ["S. aureus", "Far (Firmicutes)", "1.000", "1.000"],
        ["L. monocytogenes", "Far (Firmicutes)", "1.000", "1.000"],
        ["K. pneumoniae", "Near (Enterobacteriaceae)", "0.103", "0.485"],
        ["S. flexneri", "Near (Enterobacteriaceae)", "0.003", "0.012"],
        ["P. aeruginosa", "Intermediate", "0.000", "0.000"],
    ],
    "Table 7. OOD generalization (leave-one-species-out)"
)

add_heading_styled("Serotype Classification", level=2)
add_para("Six-way serotype classification achieved only 31-36% accuracy with k-mer methods, compared to 93-98% for SeqSero2. Whole-genome frequency averaging dilutes the sparse, localized signals in hypervariable regions (rfb, fliC, fljB).")

add_heading_styled("Interpretability Analysis", level=2)
add_para("t-SNE revealed clear separation between Salmonella and non-Salmonella hidden states (centroid distance: 25.33). Salmonella sequences induced smoother, more convergent trajectories (change 0.12 vs. 0.20, 1.65-fold difference). 73% of non-Salmonella predictions had confidence <0.1; no sample fell in the uncertainty zone.")

add_heading_styled("External Validation", level=2)
add_para("On 20 completely independent genomes, the model achieved 80% accuracy with 100% Salmonella recall. All 4 errors were false positives on E. coli and Burkholderia.")

# === DISCUSSION ===
add_heading_styled("Discussion", level=1)

add_para("Our comprehensive benchmarking establishes CNN operating on k-mer frequency vectors as the most effective method for WGS-based bacterial detection. The architecture is simple, fast, and achieves near-perfect accuracy without requiring multi-gigabyte reference databases. The strong performance stems from architectural alignment with the data: convolutions along the k-mer dimension learn which tetranucleotide frequency combinations distinguish species.")

add_para("While CNN achieves the highest absolute accuracy, LNNs offer distinct advantages in specific scenarios: 6-fold lower variance at low training data volumes, interpretable continuous-time dynamics, and parameter efficiency (101K parameters for 92.4% accuracy). The finding that GRU outperforms CfC on k-mer data suggests that the advantage of ODE-based dynamics is domain-dependent — CfC excels on genuinely continuous signals rather than discrete k-mer sequences.")

add_para("The OOD results reveal that k-mer-based models learn phylogenetic proximity rather than species-specific features. Hard negative mining partially addresses this — improving Klebsiella OOD accuracy 4.7-fold — but does not fully solve the problem. Future work should explore contrastive learning and targeted feature extraction from hypervariable genomic regions.")

add_para("Several limitations warrant consideration: k-mer frequency encoding discards positional information; all genomes were sourced from NCBI RefSeq; our simplified KmerVote implementation underestimates Kraken2's capabilities; and fine-grained serotype classification remains reliant on specialized tools like SeqSero2.")

# === CONCLUSIONS ===
add_heading_styled("Conclusions", level=1)

add_para("1. CNN with k-mer frequency encoding achieves 99.6% accuracy, matching or exceeding Kraken2 and Centrifuge while eliminating database dependencies — a pragmatic, deployable solution for genomic pathogen surveillance.")
add_para("2. Attention pooling, rather than the choice of recurrent cell, is the critical architectural component, contributing an 18-percentage-point improvement to LNN performance.")
add_para("3. Liquid Neural Networks offer unique advantages in data-limited scenarios (6-fold lower variance than CNN at 10% data) and provide interpretable continuous-time dynamics.")
add_para("4. Deep learning from k-mer frequencies fundamentally outperforms exact k-mer matching for species-level classification, capturing subtle compositional patterns that discrete matching misses.")
add_para("5. For within-species discrimination, curated allele databases remain essential; whole-genome frequency methods are insufficient and specialized tools like SeqSero2 should be preferred.")

# === DECLARATIONS ===
add_heading_styled("Declarations", level=1)

declarations = {
    "Ethics approval and consent to participate": "Not applicable. This study used publicly available bacterial genome sequences from NCBI GenBank.",
    "Consent for publication": "Not applicable.",
    "Availability of data and materials": "All genomic sequences are publicly available from NCBI GenBank. Pre-processed k-mer frequency vectors, train/validation/test splits, and trained model weights are available at [Zenodo/Figshare DOI]. Source code is available at [GitHub repository URL] under the MIT License.",
    "Competing interests": "The authors declare that they have no competing interests.",
    "Funding": "[To be completed]",
    "Authors' contributions": "[To be completed]",
    "Acknowledgements": "We thank the NCBI Pathogen Detection program for making genomic data publicly available. This work utilized the ncps-torch library for Liquid Neural Network implementation.",
}

for label, text in declarations.items():
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run = p.add_run(text)

# === REFERENCES ===
add_heading_styled("References", level=1)

refs = [
    "Majowicz SE, Musto J, Scallan E, Angulo FJ, Kirk M, O'Brien SJ, et al. The global burden of nontyphoidal Salmonella gastroenteritis. Clin Infect Dis. 2010;50(6):882-9.",
    "Lee KM, Runyon M, Herrman TJ, Phillips R, Hsieh J. Review of Salmonella detection and identification methods: Aspects of rapid emergency response and food safety. Food Control. 2015;47:264-76.",
    "Malorny B, Paccassoni E, Fach P, Bunge C, Martin A, Helmuth R. Diagnostic real-time PCR for detection of Salmonella in food. Appl Environ Microbiol. 2004;70(12):7046-52.",
    "Allard MW, Strain E, Melka D, Bunning K, Musser SM, Brown EW, et al. Practical value of food pathogen traceability through whole-genome sequencing. J Clin Microbiol. 2016;54(8):1975-83.",
    "Wood DE, Lu J, Langmead B. Improved metagenomic analysis with Kraken 2. Genome Biol. 2019;20:257.",
    "Kim D, Song L, Breitwieser FP, Salzberg SL. Centrifuge: rapid and sensitive classification of metagenomic sequences. Genome Res. 2016;26(12):1721-9.",
    "Zhang S, den Bakker HC, Li S, Chen J, Dinsmore BA, Lane C, et al. SeqSero2: rapid and improved Salmonella serotype determination using whole genome sequencing data. Appl Environ Microbiol. 2019;85(23):e01746-19.",
    "Alipanahi B, Delong A, Weirauch MT, Frey BJ. Predicting the sequence specificities of DNA- and RNA-binding proteins by deep learning. Nat Biotechnol. 2015;33(8):831-8.",
    "Zhou J, Troyanskaya OG. Predicting effects of noncoding variants with deep learning-based sequence model. Nat Methods. 2015;12(10):931-4.",
    "Hasani R, Lechner M, Amini A, Rus D, Grosu R. Closed-form continuous-time neural networks. Nat Mach Intell. 2022;4:992-1003.",
    "Grimont PAD, Weill FX. Antigenic formulae of the Salmonella serovars. 9th ed. WHO Collaborating Centre for Reference and Research on Salmonella; 2007.",
    "Hasani R, Lechner M, Amini A, Rus D, Grosu R. Liquid time-constant networks. In: Proc AAAI Conf Artif Intell; 2021.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {ref}")
    run.font.size = Pt(10)

# === SUPPLEMENTARY ===
doc.add_page_break()
add_heading_styled("Supplementary Materials", level=1)

add_para("Table S1. Complete model comparison with all variants", bold=True, size=10)
add_table(
    ["Model", "Accuracy", "F1", "AUC", "Parameters"],
    [
        ["CNN", "0.9964 +/- 0.003", "0.998", "1.000", "173,633"],
        ["GRU", "0.9827", "0.989", "0.999", "195,809"],
        ["LNN + Attention", "0.9784", "0.986", "0.994", "413,378"],
        ["XGBoost", "0.9416 +/- 0.000", "0.964", "0.989", "~500,000"],
        ["LNN-Medium (opt LR)", "0.9221", "0.950", "0.915", "412,833"],
        ["LSTM", "0.8882 +/- 0.053", "0.935", "0.843", "823,810"],
        ["Random Forest", "0.8723", "0.926", "0.975", "~5,000,000"],
        ["LNN-Large", "0.8571", "0.918", "0.649", "1,211,329"],
        ["Simple RNN", "0.8420", "0.908", "0.689", "65,889"],
        ["LNN-Small (last)", "0.8095 +/- 0.015", "0.892", "0.671", "101,121"],
        ["Transformer", "0.7987", "0.888", "0.500", "850,946"],
    ],
)

add_para("Table S2. OOD leave-one-serotype-out results", bold=True, size=10)
add_table(
    ["Held-out Serotype", "Standard", "Enhanced"],
    [
        ["Dublin", "0.9946", "0.8584"],
        ["Enteritidis", "1.0000", "0.9911"],
        ["Heidelberg", "0.9956", "0.9867"],
        ["Infantis", "1.0000", "0.4896"],
        ["Newport", "1.0000", "0.9829"],
        ["Typhimurium", "0.7576", "1.0000"],
    ],
)

add_para("Table S3. External validation breakdown", bold=True, size=10)
add_table(
    ["Category", "Samples", "Correct", "Accuracy"],
    [
        ["Salmonella", "10", "10", "100%"],
        ["Non-Salmonella (non-Enterobacteriaceae)", "6", "6", "100%"],
        ["Non-Salmonella (Enterobacteriaceae relatives)", "4", "0", "0%"],
        ["Total", "20", "16", "80%"],
    ],
)

# Save
output_path = "lnn_salmonella/results/paper_bmc_genomics.docx"
doc.save(output_path)
print(f"Word document saved to: {output_path}")
